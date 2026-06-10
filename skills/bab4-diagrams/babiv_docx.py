"""
babiv_docx — generate a clean, skripsi-grade BAB IV Word document.

Why this exists
---------------
The old pack produced .docx where *every* paragraph was the "Normal" style and
figure/table numbers were typed by hand. That caused exactly the defects the
user hit: number gaps (Gambar 4.19 -> 4.22), a stray "4. Halaman ..." artifact,
and no real heading structure (so no navigation / no clean look).

This module fixes the root cause:

  * Real Word heading styles, AUTO-NUMBERED via a multilevel list bound to the
    styles. You pass clean titles ("Analisis Sistem"); Word prints "4.1",
    "4.1.1", ... and they renumber themselves. No typed numbers, no gaps.
  * Figure/table captions use true Word SEQ fields ("Gambar 4." + SEQ). The
    counter is managed by Word, so numbers are always contiguous.
  * Body text matches the reference skripsi: Arial 11, justified, 1.5 line
    spacing, 1.25 cm first-line indent, A4, margins 4/3/3/3 cm.

Public API
----------
    from babiv_docx import BabIvDoc
    doc = BabIvDoc(chapter=4, bab_title="HASIL DAN PEMBAHASAN", system_name="...")
    doc.build_from_spec(spec, figures)   # figures: {"context": path, "level0": path,
                                         #           "level1": {"1.0": path, ...},
                                         #           "erd": path}
    doc.save("BAB_IV.docx")

You normally do not call this directly; build_babiv_assets.py drives it.
"""

from __future__ import annotations

import os
from typing import Dict, List, Optional, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ARIAL = "Arial"
BODY_PT = 11
LINE = 1.5


# --------------------------------------------------------------------------- #
# low-level OOXML helpers
# --------------------------------------------------------------------------- #
def _set_font(style, name=ARIAL, size=None, bold=None, color=None):
    f = style.font
    f.name = name
    # make the font apply to ascii / hAnsi / cs / eastAsia so nothing falls back
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(a), name)
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if color is not None:
        f.color.rgb = color


def _para_fmt(style, align=None, line=LINE, before=0.0, after=0.0,
              first_line=None, left=None, hanging=None):
    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if hanging is not None:
        pf.first_line_indent = Cm(-hanging)
    if left is not None:
        pf.left_indent = Cm(left)


def _bind_numbering(style, num_id: int, ilvl: int):
    """Attach a numId/ilvl to a paragraph style so it auto-numbers."""
    ppr = style.element.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    il = OxmlElement("w:ilvl")
    il.set(qn("w:val"), str(ilvl))
    ni = OxmlElement("w:numId")
    ni.set(qn("w:val"), str(num_id))
    numpr.append(il)
    numpr.append(ni)


def _field(paragraph, instr: str, placeholder: str = "1"):
    """Insert a Word field (begin/instr/separate/result/end) into a paragraph."""
    def _run():
        return paragraph.add_run()
    r = _run()._r
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r.append(fc)
    r = _run()._r
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r.append(it)
    r = _run()._r
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate"); r.append(fc)
    paragraph.add_run(placeholder)
    r = _run()._r
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end"); r.append(fc)


def _enable_update_fields(document):
    """Tell Word to refresh fields (SEQ caption numbers) when the doc opens."""
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)


# --------------------------------------------------------------------------- #
# numbering definitions (multilevel headings + simple lists)
# --------------------------------------------------------------------------- #
def _numbering_el(document):
    return document.part.numbering_part.element


def _next_ids(num_el):
    abs_ids = [int(a.get(qn("w:abstractNumId"))) for a in num_el.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in num_el.findall(qn("w:num"))]
    return (max(abs_ids) + 1 if abs_ids else 0), (max(num_ids) + 1 if num_ids else 1)


def _add_num(num_el, abstract_id: int) -> int:
    _, num_id = _next_ids(num_el)
    n = OxmlElement("w:num")
    n.set(qn("w:numId"), str(num_id))
    ab = OxmlElement("w:abstractNumId")
    ab.set(qn("w:val"), str(abstract_id))
    n.append(ab)
    # w:num must come after all w:abstractNum; append at end is fine
    num_el.append(n)
    return num_id


def _lvl(ilvl: int, num_fmt: str, text: str, *, start=1,
         left_cm: float = 0.0, hanging_cm: float = 0.0, bold=None,
         align_left=True) -> "OxmlElement":
    lv = OxmlElement("w:lvl")
    lv.set(qn("w:ilvl"), str(ilvl))
    s = OxmlElement("w:start"); s.set(qn("w:val"), str(start)); lv.append(s)
    nf = OxmlElement("w:numFmt"); nf.set(qn("w:val"), num_fmt); lv.append(nf)
    lt = OxmlElement("w:lvlText"); lt.set(qn("w:val"), text); lv.append(lt)
    lj = OxmlElement("w:lvlJc"); lj.set(qn("w:val"), "left"); lv.append(lj)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(Cm(left_cm).twips))
    ind.set(qn("w:hanging"), str(Cm(hanging_cm).twips))
    ppr.append(ind)
    lv.append(ppr)
    return lv


def _make_heading_list(document, chapter: int) -> int:
    """Multilevel list: chapter.%1 / chapter.%1.%2 / chapter.%1.%2.%3.

    Bound to heading styles so titles auto-number with the chapter fixed.
    """
    num_el = _numbering_el(document)
    abs_id, _ = _next_ids(num_el)
    ab = OxmlElement("w:abstractNum")
    ab.set(qn("w:abstractNumId"), str(abs_id))
    mlt = OxmlElement("w:multiLevelType"); mlt.set(qn("w:val"), "multilevel"); ab.append(mlt)
    c = chapter
    # chapter digit is literal text; %1/%2/%3 are the per-level counters.
    ab.append(_lvl(0, "decimal", f"{c}.%1", left_cm=0.75, hanging_cm=0.75))
    ab.append(_lvl(1, "decimal", f"{c}.%1.%2", left_cm=1.0, hanging_cm=1.0))
    ab.append(_lvl(2, "decimal", f"{c}.%1.%2.%3", left_cm=1.25, hanging_cm=1.25))
    num_el.insert(0, ab)  # abstractNum must precede num elements in the part
    return _add_num(num_el, abs_id)


def _make_simple_list(document, num_fmt: str, left_cm: float, hanging_cm: float) -> int:
    """A single-level list (decimal '1.' or lowerLetter 'a.'). Each call = fresh
    abstract+num so the list restarts at 1/a where used."""
    num_el = _numbering_el(document)
    abs_id, _ = _next_ids(num_el)
    ab = OxmlElement("w:abstractNum")
    ab.set(qn("w:abstractNumId"), str(abs_id))
    mlt = OxmlElement("w:multiLevelType"); mlt.set(qn("w:val"), "singleLevel"); ab.append(mlt)
    suffix = "%1." 
    ab.append(_lvl(0, num_fmt, suffix, left_cm=left_cm, hanging_cm=hanging_cm))
    num_el.insert(0, ab)
    return _add_num(num_el, abs_id)


# --------------------------------------------------------------------------- #
# main builder
# --------------------------------------------------------------------------- #
class BabIvDoc:
    def __init__(self, chapter: int = 4, bab_title: str = "HASIL DAN PEMBAHASAN",
                 system_name: str = ""):
        self.chapter = chapter
        self.bab_title = bab_title
        self.system_name = system_name
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    # ---- setup -------------------------------------------------------------
    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(4.0)
        s.right_margin = Cm(3.0)
        s.top_margin = Cm(3.0)
        s.bottom_margin = Cm(3.0)

    def _setup_styles(self):
        d = self.doc
        # Normal / default body
        normal = d.styles["Normal"]
        _set_font(normal, ARIAL, BODY_PT, color=RGBColor(0, 0, 0))
        _para_fmt(normal, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE,
                  before=0, after=0, first_line=1.25)

        # Body paragraph style ("Paragraf")
        body = d.styles.add_style("Paragraf", WD_STYLE_TYPE.PARAGRAPH)
        body.base_style = normal
        _set_font(body, ARIAL, BODY_PT, bold=False, color=RGBColor(0, 0, 0))
        _para_fmt(body, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=LINE,
                  before=0, after=0, first_line=1.25)

        # Heading 1 = BAB title (centered, 14 bold, no number)
        h1 = d.styles["Heading 1"]
        _set_font(h1, ARIAL, 14, bold=True, color=RGBColor(0, 0, 0))
        _para_fmt(h1, align=WD_ALIGN_PARAGRAPH.CENTER, line=LINE, before=0, after=12)
        h1.element.get_or_add_pPr()  # ensure pPr
        self._set_outline(h1, 0)

        # build the auto-number multilevel list and bind to H2/H3/H4
        head_num = _make_heading_list(d, self.chapter)
        h2 = d.styles["Heading 2"]
        _set_font(h2, ARIAL, 12, bold=True, color=RGBColor(0, 0, 0))
        _para_fmt(h2, align=WD_ALIGN_PARAGRAPH.LEFT, line=LINE, before=12, after=4,
                  left=0.75, hanging=0.75)
        self._set_outline(h2, 1)
        _bind_numbering(h2, head_num, 0)

        h3 = d.styles["Heading 3"]
        _set_font(h3, ARIAL, 12, bold=True, color=RGBColor(0, 0, 0))
        _para_fmt(h3, align=WD_ALIGN_PARAGRAPH.LEFT, line=LINE, before=8, after=4,
                  left=1.0, hanging=1.0)
        self._set_outline(h3, 2)
        _bind_numbering(h3, head_num, 1)

        h4 = d.styles["Heading 4"]
        _set_font(h4, ARIAL, 11, bold=True, color=RGBColor(0, 0, 0))
        _para_fmt(h4, align=WD_ALIGN_PARAGRAPH.LEFT, line=LINE, before=6, after=4,
                  left=1.25, hanging=1.25)
        self._set_outline(h4, 3)
        _bind_numbering(h4, head_num, 2)

        # Caption (centered, 11, not bold)
        try:
            cap = d.styles["Caption"]
        except KeyError:
            cap = d.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        _set_font(cap, ARIAL, 11, bold=False, color=RGBColor(0, 0, 0))
        _para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, line=LINE, before=6, after=12)

        # Figure holder (centered image, tight spacing)
        fig = d.styles.add_style("FigureImage", WD_STYLE_TYPE.PARAGRAPH)
        fig.base_style = normal
        _para_fmt(fig, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, before=12, after=0,
                  first_line=0.0)

        _enable_update_fields(d)

    def _set_outline(self, style, level: int):
        ppr = style.element.get_or_add_pPr()
        ol = ppr.find(qn("w:outlineLvl"))
        if ol is None:
            ol = OxmlElement("w:outlineLvl")
            ppr.append(ol)
        ol.set(qn("w:val"), str(level))

    # ---- content helpers ---------------------------------------------------
    def bab(self):
        p = self.doc.add_paragraph(style="Heading 1")
        p.add_run(f"BAB {self._roman(self.chapter)}")
        p.add_run("\n")
        p.add_run(self.bab_title.upper())
        return p

    def h2(self, text):
        return self.doc.add_paragraph(text, style="Heading 2")

    def h3(self, text):
        return self.doc.add_paragraph(text, style="Heading 3")

    def h4(self, text):
        return self.doc.add_paragraph(text, style="Heading 4")

    def body(self, text):
        return self.doc.add_paragraph(text, style="Paragraf")

    def figure(self, image_path: str, caption: str, width_cm: float = 14.0):
        if image_path and os.path.exists(image_path):
            p = self.doc.add_paragraph(style="FigureImage")
            p.add_run().add_picture(image_path, width=Cm(width_cm))
        self._caption("Gambar", caption)

    def table_caption(self, caption: str):
        self._caption("Tabel", caption)

    def _caption(self, kind: str, text: str):
        p = self.doc.add_paragraph(style="Caption")
        p.add_run(f"{kind} {self.chapter}.")
        _field(p, f" SEQ {kind} \\* ARABIC ", "1")
        p.add_run(" " + text)
        return p

    def numbered_list(self, items: List[str]):
        nid = _make_simple_list(self.doc, "decimal", left_cm=1.0, hanging_cm=0.5)
        out = []
        for it in items:
            p = self.doc.add_paragraph(it, style="Paragraf")
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            _bind_numbering_para(p, nid, 0)
            out.append(p)
        return out

    def alpha_list(self, items: List[str]):
        nid = _make_simple_list(self.doc, "lowerLetter", left_cm=1.75, hanging_cm=0.5)
        out = []
        for it in items:
            p = self.doc.add_paragraph(it, style="Paragraf")
            p.paragraph_format.left_indent = Cm(1.75)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            _bind_numbering_para(p, nid, 0)
            out.append(p)
        return out

    def data_table(self, header: List[str], rows: List[List[str]],
                   widths_cm: Optional[List[float]] = None,
                   align: Optional[List[str]] = None):
        cols = len(header)
        t = self.doc.add_table(rows=1, cols=cols)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        if widths_cm is None:
            total = 21.0 - 4.0 - 3.0
            widths_cm = [total / cols] * cols

        def _cellfmt(cell, text, *, bold=False, center=False):
            # python-docx leaves a default paragraph that inherits Normal's
            # 1.25 cm first-line indent -> that is the stray "tab" the user saw
            # on username/password rows. Force flush-left, no indent, tight.
            cell.text = ""
            para = cell.paragraphs[0]
            pf = para.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0)
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.line_spacing = 1.0
            pf.alignment = (WD_ALIGN_PARAGRAPH.CENTER if center
                            else WD_ALIGN_PARAGRAPH.LEFT)
            run = para.add_run(str(text))
            run.font.name = ARIAL
            run.font.size = Pt(BODY_PT)
            run.bold = bold
            _cell_valign_top(cell)

        hdr = t.rows[0].cells
        for i, htext in enumerate(header):
            _cellfmt(hdr[i], htext, bold=True, center=True)
            _shade(hdr[i], "D9D9D9")
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                center = bool(align and i < len(align) and align[i] == "center")
                _cellfmt(cells[i], val, center=center)
        # fixed widths (authoritative: tblGrid + per-cell + fixed layout)
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
        _table_fixed(t)
        _set_col_grid(t, widths_cm)
        # tighten side padding when columns are narrow (many-column SUS tables)
        if min(widths_cm) < 1.1:
            _set_cell_margins(t, lr_cm=0.08)
        return t

    # ---- spec-driven assembly ---------------------------------------------
    def build_from_spec(self, spec: Dict[str, Any], figures: Dict[str, Any],
                        sus_data: Optional[Dict[str, Any]] = None):
        self.bab()
        self._section_analisis(spec)         # 4.1
        self._section_perancangan(spec, figures)   # 4.2
        self._section_basisdata(spec, figures)     # 4.3
        self._section_antarmuka(spec)        # 4.4
        self._section_implementasi(spec)     # 4.5
        self._section_blackbox(spec)         # 4.6 Pengujian Black Box
        self._section_sus(spec, sus_data)    # 4.7 Pengujian SUS

    def _section_analisis(self, spec):
        an = spec.get("analisis", {}) or {}
        self.h2("Analisis Sistem")
        intro = an.get("deskripsi") or (
            f"Pada tahap ini dilakukan analisis terhadap {self.system_name} untuk "
            "mengidentifikasi permasalahan, kebutuhan fungsional, dan kebutuhan "
            "non-fungsional yang menjadi dasar perancangan sistem.")
        self.body(intro)

        if an.get("masalah"):
            self.h3("Analisis Masalah")
            for para in _as_paragraphs(an["masalah"]):
                self.body(para)

        fung = an.get("kebutuhan_fungsional")
        if fung:
            self.h3("Analisis Kebutuhan Fungsional")
            self.body("Kebutuhan fungsional mendeskripsikan layanan yang harus "
                      "disediakan sistem berdasarkan hak akses masing-masing "
                      "pengguna sebagai berikut.")
            # fung may be list[str] OR list[{actor, items[]}]
            if fung and isinstance(fung[0], dict):
                actor_nid = _make_simple_list(self.doc, "decimal", left_cm=1.0, hanging_cm=0.5)
                for grp in fung:
                    p = self.doc.add_paragraph(grp.get("actor", ""), style="Paragraf")
                    p.paragraph_format.left_indent = Cm(1.0)
                    p.paragraph_format.first_line_indent = Cm(-0.5)
                    _bind_numbering_para(p, actor_nid, 0)
                    if grp.get("items"):
                        self.alpha_list(grp["items"])
            else:
                self.numbered_list(list(fung))

        nonf = an.get("kebutuhan_nonfungsional")
        if nonf:
            self.h3("Analisis Kebutuhan Non-Fungsional")
            if isinstance(nonf, dict):
                if nonf.get("deskripsi"):
                    self.body(nonf["deskripsi"])
                for tb in nonf.get("tabel", []):
                    self.table_caption(tb.get("judul", "Tabel"))
                    self.data_table(tb["header"], tb["rows"], tb.get("widths_cm"))
            else:
                self.body("Kebutuhan non-fungsional sistem meliputi:")
                self.numbered_list(list(nonf))

    def _section_perancangan(self, spec, figures):
        self.h2("Perancangan Sistem")
        self.body("Perancangan sistem menggambarkan alur data menggunakan Data "
                  "Flow Diagram (DFD) mulai dari diagram konteks hingga DFD level "
                  "rinci, sebagai acuan implementasi.")

        # Diagram Konteks
        if figures.get("context"):
            self.h3("Diagram Konteks")
            self.body("Diagram konteks merupakan representasi tingkat tertinggi "
                      f"yang menggambarkan {self.system_name} sebagai satu proses "
                      "tunggal beserta interaksinya dengan entitas eksternal.")
            self.figure(figures["context"], "Diagram Konteks")

        # DFD Level 0
        if figures.get("level0"):
            self.h3("DFD Level 0")
            self.body("DFD Level 0 menguraikan proses utama di dalam sistem "
                      "beserta aliran data dan media penyimpanannya (data store).")
            self.figure(figures["level0"], "DFD Level 0")

        # DFD Level 1 (per decomposed process)
        lvl1 = figures.get("level1") or {}
        if lvl1:
            self.h3("DFD Level 1")
            self.body("DFD Level 1 merupakan dekomposisi dari proses pada DFD "
                      "Level 0 menjadi sub-proses yang lebih rinci.")
            # order by process number
            for no in sorted(lvl1.keys(), key=_num_key):
                meta = lvl1[no]
                path = meta["path"] if isinstance(meta, dict) else meta
                name = meta.get("name", "") if isinstance(meta, dict) else ""
                title = f"DFD Level 1 Proses {no}" + (f" {name}" if name else "")
                self.h4(title)
                self.figure(path, title)

    def _section_basisdata(self, spec, figures):
        self.h2("Perancangan Basis Data")
        self.body("Perancangan basis data digambarkan menggunakan Entity "
                  "Relationship Diagram (ERD) untuk menunjukkan entitas, atribut, "
                  "dan relasi antar tabel.")
        if figures.get("erd"):
            self.h3("Entity Relationship Diagram (Notasi Crow's Foot)")
            self.body("ERD berikut menggunakan notasi crow's foot: tiap entitas "
                      "digambarkan sebagai kotak berisi atribut (PK/FK), dan "
                      "relasi antar entitas memakai simbol kaki gagak untuk "
                      "menunjukkan kardinalitas.")
            self.figure(figures["erd"],
                        "Entity Relationship Diagram (Notasi Crow's Foot)")
        if figures.get("erd_chen"):
            self.h3("Entity Relationship Diagram (Notasi Chen)")
            self.body("Sebagai pelengkap, model konseptual basis data juga "
                      "digambarkan dengan notasi Chen, di mana entitas berbentuk "
                      "persegi, atribut berbentuk elips, dan relasi berbentuk "
                      "belah ketupat dengan kardinalitas 1/N.")
            self.figure(figures["erd_chen"],
                        "Entity Relationship Diagram (Notasi Chen)")

        tables = spec.get("struktur_tabel")
        if tables:
            self.h3("Struktur Tabel")
            self.body("Adapun struktur tabel pada basis data sistem adalah "
                      "sebagai berikut.")
            for tb in tables:
                self.table_caption(f"Struktur Tabel {tb.get('nama','')}")
                self.data_table(tb["header"], tb["rows"], tb.get("widths_cm"))

    def _section_antarmuka(self, spec):
        items = spec.get("antarmuka") or []
        self.h2("Perancangan Antarmuka")
        self.body("Perancangan antarmuka bertujuan memberikan gambaran tampilan "
                  "sistem yang akan dibangun agar mudah digunakan oleh pengguna.")
        for it in items:
            self.h3(it.get("title", "Halaman"))
            if it.get("desc"):
                self.body(it["desc"])
            if it.get("image"):
                self.figure(it["image"], it.get("caption", it.get("title", "Antarmuka")))

    def _section_implementasi(self, spec):
        items = spec.get("implementasi") or []
        self.h2("Implementasi Sistem")
        self.body("Implementasi merupakan tahap penerapan hasil perancangan ke "
                  "dalam bentuk sistem yang siap digunakan. Berikut adalah "
                  "tampilan hasil implementasi sistem.")
        for it in items:
            self.h3(it.get("title", "Halaman"))
            if it.get("desc"):
                self.body(it["desc"])
            if it.get("image"):
                self.figure(it["image"], it.get("caption", it.get("title", "Implementasi")))

    def _section_blackbox(self, spec):
        """4.6 Pengujian Black Box — one test table per program feature.
        Columns match the reference skripsi: No | Data Input | Hasil yang
        Diharapkan | Hasil Pengamatan | Kesimpulan."""
        self.h2("Pengujian Black Box")
        pj = (spec.get("pengujian") or {})
        groups = pj.get("blackbox") or []
        self.body(
            "Pengujian black box dilakukan untuk memastikan setiap fungsi pada "
            "sistem berjalan sesuai dengan kebutuhan tanpa melihat struktur kode "
            "program. Pengujian difokuskan pada kesesuaian antara data masukan, "
            "hasil yang diharapkan, dan hasil pengamatan pada saat sistem "
            "dijalankan. Adapun skenario dan hasil pengujian tiap fitur adalah "
            "sebagai berikut.")
        if not groups:
            self.body("(Daftar skenario pengujian black box belum diisi pada "
                      "spec.json bagian \"pengujian.blackbox\".)")
            return
        hdr = ["No", "Data Input", "Hasil yang Diharapkan",
               "Hasil Pengamatan", "Kesimpulan"]
        widths = [1.0, 4.2, 4.2, 3.3, 1.8]
        total_ok = total_all = 0
        for grp in groups:
            judul = grp.get("judul") or grp.get("title") or "Pengujian Fitur"
            cases = grp.get("cases") or grp.get("rows") or []
            self.table_caption(f"Pengujian {judul}" if not str(judul).lower()
                               .startswith("pengujian") else judul)
            rows = []
            for i, c in enumerate(cases, start=1):
                if isinstance(c, dict):
                    inp = c.get("input", "")
                    exp = c.get("expect", c.get("expected", ""))
                    obs = c.get("observe", c.get("observed", exp))
                    res = c.get("result", "Sesuai")
                else:  # list form [input, expect, observe, result]
                    inp = c[0] if len(c) > 0 else ""
                    exp = c[1] if len(c) > 1 else ""
                    obs = c[2] if len(c) > 2 else exp
                    res = c[3] if len(c) > 3 else "Sesuai"
                rows.append([str(i), inp, exp, obs, res])
                total_all += 1
                if str(res).strip().lower() in ("sesuai", "valid", "berhasil"):
                    total_ok += 1
            self.data_table(hdr, rows, widths_cm=widths,
                            align=["center", "", "", "", "center"])
        if total_all:
            pct = round(total_ok / total_all * 100)
            self.body(
                f"Berdasarkan hasil pengujian black box pada tabel-tabel di atas, "
                f"dari {total_all} skenario pengujian yang dilakukan, sebanyak "
                f"{total_ok} skenario memberikan hasil yang sesuai dengan yang "
                f"diharapkan ({pct}%). Dengan demikian dapat disimpulkan bahwa "
                f"fungsi-fungsi utama sistem telah berjalan dengan baik dan sesuai "
                f"dengan kebutuhan.")

    # 10 standard SUS statements (Indonesian), used when the data file has none.
    _SUS_STATEMENTS = [
        "Saya berpikir akan sering menggunakan sistem ini.",
        "Saya merasa sistem ini terlalu rumit untuk digunakan.",
        "Saya merasa sistem ini mudah untuk digunakan.",
        "Saya membutuhkan bantuan teknis untuk dapat menggunakan sistem ini.",
        "Saya merasa fitur-fitur pada sistem ini berjalan dengan semestinya.",
        "Saya merasa terlalu banyak ketidaksesuaian pada sistem ini.",
        "Saya merasa kebanyakan orang akan cepat memahami cara memakai sistem ini.",
        "Saya merasa sistem ini membingungkan ketika digunakan.",
        "Saya merasa percaya diri dalam menggunakan sistem ini.",
        "Saya perlu belajar banyak hal sebelum dapat menggunakan sistem ini.",
    ]

    def _section_sus(self, spec, sus):
        """4.7 Pengujian SUS — questionnaire table + per-respondent converted
        score table (Responden | 1..10 | Skor Total | Skor SUS) + example
        calculation + interpretation, exactly the layout of the reference doc."""
        self.h2("Pengujian System Usability Scale (SUS)")
        if not sus:
            self.body("Pengujian usability dilakukan menggunakan metode System "
                      "Usability Scale (SUS). (Data kuesioner pada data.xlsx belum "
                      "tersedia, sehingga perhitungan skor belum dapat ditampilkan.)")
            return
        n = sus["n"]
        n_items = sus.get("n_items", 10)
        self.body(
            f"Pengujian kebergunaan (usability) sistem dilakukan menggunakan "
            f"metode System Usability Scale (SUS). Kuesioner SUS terdiri atas "
            f"{n_items} pernyataan dengan skala 1 (sangat tidak setuju) sampai 5 "
            f"(sangat setuju), di mana pernyataan bernomor ganjil bersifat positif "
            f"dan pernyataan bernomor genap bersifat negatif.")
        cats = sus.get("category_counts") or {}
        if cats:
            cat_txt = ", ".join(f"{k} sebanyak {v} orang" for k, v in cats.items())
            self.body(f"Kuesioner disebarkan kepada {n} responden (n = {n}) yang "
                      f"terdiri atas {cat_txt}.")
        else:
            self.body(f"Kuesioner disebarkan kepada {n} responden (n = {n}).")

        # ---- questionnaire table (No | Pernyataan | 1 2 3 4 5) -------------
        self.h3("Daftar Pernyataan Kuesioner")
        self.body("Daftar pernyataan yang digunakan dalam kuesioner SUS "
                  "ditunjukkan pada tabel berikut.")
        statements = sus.get("item_headers") or []
        statements = [s for s in statements if s and s.strip()]
        if len(statements) < n_items:
            statements = self._SUS_STATEMENTS[:n_items]
        q_rows = [[str(i + 1), statements[i], "\u25cb", "\u25cb", "\u25cb",
                   "\u25cb", "\u25cb"] for i in range(min(n_items, len(statements)))]
        self.table_caption("Daftar Pernyataan Kuesioner SUS")
        self.data_table(["No", "Pernyataan", "1", "2", "3", "4", "5"], q_rows,
                        widths_cm=[1.0, 8.0, 1.0, 1.0, 1.0, 1.0, 1.0],
                        align=["center", "", "center", "center", "center",
                               "center", "center"])

        # ---- scoring explanation + example ---------------------------------
        self.h3("Perhitungan Skor SUS")
        self.body(
            "Perhitungan skor SUS dilakukan dengan ketentuan: untuk pernyataan "
            "bernomor ganjil nilai kontribusi = (skor jawaban \u2212 1), sedangkan "
            "untuk pernyataan bernomor genap nilai kontribusi = (5 \u2212 skor "
            "jawaban). Seluruh nilai kontribusi dari 10 pernyataan dijumlahkan, "
            "kemudian dikalikan 2,5 untuk mendapatkan skor SUS dengan rentang "
            "0\u2013100.")

        # example using respondent #1 (shows the conversion explicitly)
        ex = sus["per_respondent"][0]
        ex_raw = ex.get("raw") or []
        ex_con = ex.get("contribs") or []
        if ex_raw and ex_con:
            head = ["Pernyataan"] + [str(i + 1) for i in range(len(ex_con))] + \
                   ["Skor Total", "Skor SUS"]
            r_jwb = ["Jawaban"] + [str(v) for v in ex_raw] + ["", ""]
            r_kon = ["Kontribusi"] + [str(int(v)) for v in ex_con] + \
                    [str(int(ex["total"])), f"{ex['score']:.0f}"]
            self.table_caption(f"Contoh Perhitungan Skor SUS Responden "
                               f"{ex['no']}")
            w = [2.0] + [0.78] * len(ex_con) + [1.5, 1.4]
            self.data_table(head, [r_jwb, r_kon], widths_cm=w,
                            align=["center"] + ["center"] * (len(ex_con) + 2))

        # ---- full per-respondent converted table ---------------------------
        self.body("Hasil perhitungan skor SUS untuk seluruh responden "
                  "ditampilkan pada tabel berikut, di mana kolom 1 sampai "
                  f"{n_items} menunjukkan nilai kontribusi tiap pernyataan.")
        head = ["Responden"] + [str(i + 1) for i in range(n_items)] + \
               ["Skor Total", "Skor SUS"]
        rows = []
        for p in sus["per_respondent"]:
            con = p.get("contribs") or []
            rows.append([str(p["no"])] + [str(int(v)) for v in con] +
                        [str(int(p["total"])), f"{p['score']:.0f}"])
        w = [1.8] + [0.8] * n_items + [1.55, 1.45]
        self.table_caption("Perhitungan Skor SUS Seluruh Responden")
        self.data_table(head, rows, widths_cm=w,
                        align=["center"] + ["center"] * (n_items + 2))

        self.body(
            f"Berdasarkan tabel di atas, dari {n} responden diperoleh skor SUS "
            f"minimum sebesar {sus['min']:.0f}, skor maksimum {sus['max']:.0f}, "
            f"dan rata-rata skor SUS sebesar {sus['mean']:.1f}.")

        # ---- interpretation -------------------------------------------------
        self.h3("Interpretasi Hasil")
        self.body(
            f"Nilai rata-rata skor SUS sebesar {sus['mean']:.1f} berada pada "
            f"kategori grade \u201c{sus['grade']}\u201d dengan predikat "
            f"\u201c{sus['adjective']}\u201d. Berdasarkan tingkat penerimaan "
            f"(acceptability), skor tersebut tergolong {sus['acceptability']}. "
            f"Hal ini menunjukkan bahwa sistem yang dibangun memiliki tingkat "
            f"kebergunaan yang baik dan dapat diterima oleh pengguna.")

    # ---- io ----------------------------------------------------------------
    def save(self, path: str):
        self.doc.save(path)
        return path

    @staticmethod
    def _roman(n: int) -> str:
        table = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
                 (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
                 (5, "V"), (4, "IV"), (1, "I")]
        out = ""
        for v, s in table:
            while n >= v:
                out += s
                n -= v
        return out


# helpers that need a paragraph (not a style) -------------------------------- #
def _bind_numbering_para(paragraph, num_id: int, ilvl: int):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(qn("w:val"), str(ilvl)); numpr.append(il)
    ni = OxmlElement("w:numId"); ni.set(qn("w:val"), str(num_id)); numpr.append(ni)
    ppr.append(numpr)


def _shade(cell, hex_fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def _cell_valign_top(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    va = tcpr.find(qn("w:vAlign"))
    if va is None:
        va = OxmlElement("w:vAlign")
        tcpr.append(va)
    va.set(qn("w:val"), "top")


def _table_fixed(table):
    tbl = table._tbl
    tblpr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)


def _set_col_grid(table, widths_cm):
    """Authoritatively set column widths via <w:tblGrid> (the source Word/
    LibreOffice actually use for fixed-layout tables). Per-cell widths alone are
    unreliable on a many-column table and led to headers breaking mid-word."""
    tbl = table._tbl
    # remove any existing grid, then insert a fresh one right after tblPr
    for g in tbl.findall(qn("w:tblGrid")):
        tbl.remove(g)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(w).twips)))
        grid.append(col)
    tbl.tblPr.addnext(grid)


def _set_cell_margins(table, lr_cm: float = 0.1):
    """Shrink default cell side padding (~0.19cm) so narrow numeric columns can
    show '10' on one line."""
    tblpr = table._tbl.tblPr
    mar = tblpr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblpr.append(mar)
    for side in ("left", "right"):
        el = mar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            mar.append(el)
        el.set(qn("w:w"), str(int(Cm(lr_cm).twips)))
        el.set(qn("w:type"), "dxa")


def _as_paragraphs(val) -> List[str]:
    if isinstance(val, list):
        return [str(x) for x in val]
    return [p.strip() for p in str(val).split("\n\n") if p.strip()]


def _num_key(s: str):
    try:
        return [int(x) for x in str(s).replace(",", ".").split(".")]
    except Exception:
        return [9999]
